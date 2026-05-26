from flask import Flask, render_template, request, send_file
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import uuid
from openai import OpenAI
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

app = Flask(__name__)

# For Vercel serverless environment, use /tmp for writable files
app.config['UPLOAD_FOLDER'] = '/tmp/uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB limit

# Ensure upload directory exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Initialize DeepSeek client
client = OpenAI(
    api_key=os.getenv("DEEPSEEK_API_KEY"),
    base_url="https://api.deepseek.com"
)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/generate', methods=['POST'])
def generate():
    # Get form data
    topic = request.form.get('topic', '')
    word_count = int(request.form.get('word_count', 1000))
    structure = request.form.get('structure', '')
    has_illustration = request.form.get('has_illustration') == 'on'
    
    # Handle image upload
    image_path = None
    if has_illustration:
        if 'illustration' in request.files:
            file = request.files['illustration']
            if file.filename != '':
                filename = str(uuid.uuid4()) + os.path.splitext(file.filename)[1]
                image_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(image_path)
    
    # Use DeepSeek AI to generate essay content (with fallback)
    essay_content = generate_essay_with_deepseek(topic, word_count, structure, has_illustration)
    
    # Create Word document
    doc_path = create_document(essay_content, topic, image_path, has_illustration)
    
    # Return Word document for download
    return send_file(doc_path, as_attachment=True, download_name=f'{topic}_课程论文.docx')

def generate_essay_with_deepseek(topic, word_count, structure, has_illustration=False):
    """
    使用DeepSeek API生成论文内容
    """
    # 构建系统提示词和用户提示词
    system_prompt = """你是一位经验丰富的大学导师，擅长指导文科本科生完成课程论文写作。
    你的任务是根据用户提供的主题、字数要求和结构偏好，生成一篇符合学术规范的课程论文。
    生成的内容应当：
    1. 语言正式、学术，但不过于晦涩
    2. 结构清晰，逻辑连贯
    3. 观点有依据，避免无根据的断言
    4. 符合"够交差型"标准 - 重点在完成度和规范性而非创新性
    5. 不要使用第一人称，保持客观中立的学术语气
    6. 字数要接近用户指定的目标范围（±10%可接受）"""
    
    # 根据结构类型构建不同的提示
    structure_desc = "三部分结构：引言（述背景提出问题）- 论点一（理论分析）- 论点二（现实案例）- 结论（综合总结提出建议）" \
                    if structure == 'three-part' else \
                    "自由结构：可包括引言、理论综述、问题分析、案例研究、结论等部分，逻辑自然连贯"
    
    illustration_note = "建议在适当位置留出1-2个插入图片的空白，图片应与论证内容相关。" \
                       if has_illustration else "纯文本论文，无需考虑插图位置。"
    
    user_prompt = f"""请为以下课程论文要求生成内容：

论文主题：{topic}
目标字数：{word_count}字左右
论文结构：{structure_desc}
其他要求：{illustration_note}

请直接返回论文正文内容，不要包含任何说明性文字、标题或格式标记。
内容应从论文正文开始（不需要重复标题），可以自然包含段落划分。"""

    try:
        # 调用DeepSeek API
        response = client.chat.completions.create(
            model="deepseek-chat",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.7,  # 平衡创造性和一致性
            max_tokens=min(word_count * 2, 2000),  # 粗略估算：中文约1字=0.5token
            stream=False
        )
        
        # 提取生成的内容
        generated_text = response.choices[0].message.content.strip()
        
        # 后处理：确保基本质量
        return post_process_essay(generated_text, topic, word_count)
        
    except Exception as e:
        # 记录错误（在实际应用中应使用日志系统）
        print(f"DeepSeek API调用失败: {str(e)}")
        # 降级到原有模板生成方式
        return generate_essay_fallback(topic, word_count, structure)

def post_process_essay(text, topic, target_word_count):
    """后处理AI生成的论文，确保基本质量"""
    # 确保内容不为空
    if not text or len(text.strip()) < 50:
        return generate_essay_fallback(topic, target_word_count, 'three-part')
    
    # 简单的字数调整：如果太短则补充结论部分
    current_words = len(text)
    if current_words < target_word_count * 0.8:
        conclusion = f"\n\n总之，{topic}是一个重要的研究课题。通过本文的分析可以看出，它具有重要的理论价值和现实意义。未来的研究可以进一步深化探讨其在不同背景下的应用。"
        text += conclusion
    # 如果太长则截取前半部分（简单处理）
    elif current_words > target_word_count * 1.2:
        # 按句子截断，尽量保持完整性
        sentences = text.split('。')
        target_sentences = max(3, int(len(sentences) * target_word_count / current_words))
        text = '。'.join(sentences[:target_sentences]) + '。'
    
    # 确保有合理的开头结构
    if not any(text.startswith(prefix) for prefix in ['引言', '第一章', '本文']):
        text = f"引言：本文将围绕\"{topic}\"展开研究。\n\n" + text
    
    return text.strip()

def generate_essay_fallback(topic, word_count, structure):
    """原有的模板生成函数作为降级方案"""
    paragraphs = []
    if structure == 'three-part':
        paragraphs.append(f"引言：本文将围绕\"{topic}\"展开讨论。")
        paragraphs.append(f"论点一：从历史角度看，{topic}具有重要意义。")
        paragraphs.append(f"论点二：从现实角度看，{topic}面临着新的挑战和机遇。")
        paragraphs.append(f"结论：综上所述，{topic}是一个值得深入研究的课题。")
    else:  # free structure
        paragraphs.append(f"随着社会的发展，{topic}问题日益突出。")
        paragraphs.append("学者们对此进行了广泛的研究。")
        paragraphs.append(f"本文认为，需要从多个维度来理解{topic}。")
        paragraphs.append(f"总之，{topic}对社会发展有着重要影响。")
    
    full_text = '。'.join(paragraphs) + '。'
    # 简单调整以接近目标字数
    if len(full_text) < word_count * 0.5:
        full_text += f" 从理论与实践相结合的视角出发，对{topic}进行系统分析有助于深化我们的认识。"
    
    return full_text

def create_document(content, topic, image_path, has_illustration):
    document = Document()
    
    # Set font and size for normal text
    style = document.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    
    # Add title
    title = document.add_heading('《{}》课程论文'.format(topic), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add content
    paragraphs = content.split('。')
    for para in paragraphs:
        if para.strip():
            p = document.add_paragraph(para.strip() + '。')
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Add image if provided and has_illustration is True
    if has_illustration and image_path and os.path.exists(image_path):
        # Add a paragraph for the image
        last_para = document.add_paragraph()
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = last_para.add_run()
        run.add_picture(image_path, width=Inches(5))
        # Add caption
        caption = document.add_paragraph('图片：{}示意图'.format(topic))
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.style = 'Caption'
    
    # Add table of contents (simplified)
    document.add_page_break()
    toc_heading = document.add_heading('目录', level=1)
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph('一、引言........................................1')
    document.add_paragraph('二、论点一......................................2')
    document.add_paragraph('三、论点二......................................3')
    document.add_paragraph('四、结论........................................4')
    
    # Add references
    document.add_page_break()
    ref_heading = document.add_heading('参考文献', level=1)
    ref_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph('[1] 某某某. 某某研究[J]. 某某期刊, 2020, 10(1): 1-10.')
    document.add_paragraph('[2] 某某某. 某某分析[M]. 某某出版社, 2019.')
    
    # Save document to a temporary location for Vercel
    filename = str(uuid.uuid4()) + '.docx'
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    document.save(filepath)
    return filepath

if __name__ == '__main__':
    app.run(debug=True)