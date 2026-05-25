from flask import Flask, render_template, request, send_file
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import uuid

app = Flask(__name__)

# For Vercel serverless environment, use /tmp for writable files
app.config['UPLOAD_FOLDER'] = '/tmp/uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB limit

# Ensure upload directory exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/generate', methods=['POST'])
def generate():
    # Get form data
    topic = request.form.get('topic', '')
    word_count = int(request.form.get('word_count', 1000))
    structure = request.form.get('structure', '')
    has_illustration = request.form.get('has_illustration') == 'on'
    
    # Handle image upload
    image_path = None
    if has_illustration:
        if 'illustration' in request.files:
            file = request.files['illustration']
            if file.filename != '':
                filename = str(uuid.uuid4()) + os.path.splitext(file.filename)[1]
                image_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(image_path)
        else:
            # Generate a placeholder image path (in reality, we might generate an image)
            # For simplicity, we'll use a placeholder or skip
            image_path = None  # We'll handle missing image in document generation
    
    # Generate essay content (simplified)
    essay_content = generate_essay(topic, word_count, structure)
    
    # Create Word document
    doc_path = create_document(essay_content, topic, image_path, has_illustration)
    
    # Return Word document for download (PDF conversion removed for Vercel compatibility)
    return send_file(doc_path, as_attachment=True, download_name=f'{topic}_课程论文.docx')

def generate_essay(topic, word_count, structure):
    # This is a placeholder for actual AI generation
    # We'll generate a simple essay based on the structure
    paragraphs = []
    if structure == 'three-part':
        paragraphs.append("引言：本文将围绕\"{}\"展开讨论。".format(topic))
        paragraphs.append("论点一：从历史角度看，{}具有重要意义。".format(topic))
        paragraphs.append("论点二：从现实角度看，{}面临着新的挑战和机遇。".format(topic))
        paragraphs.append("结论：综上所述，{}是一个值得深入研究的课题。".format(topic))
    else:  # free structure
        paragraphs.append("随着社会的发展，{}问题日益突出。".format(topic))
        paragraphs.append("学者们对此进行了广泛的研究。")
        paragraphs.append("本文认为，需要从多个维度来理解{}。".format(topic))
        paragraphs.append("总之，{}对社会发展有着重要影响。".format(topic))
    
    # Adjust to roughly meet word count (very rough approximation)
    full_text = ' '.join(paragraphs)
    # We'll just return the generated text; in reality, we'd adjust length
    return full_text

def create_document(content, topic, image_path, has_illustration):
    document = Document()
    
    # Set font and size for normal text
    style = document.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    
    # Add title
    title = document.add_heading('《{}》课程论文'.format(topic), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add content
    paragraphs = content.split('。')
    for para in paragraphs:
        if para.strip():
            p = document.add_paragraph(para.strip() + '。')
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Add image if provided and has_illustration is True
    if has_illustration and image_path and os.path.exists(image_path):
        # Add a paragraph for the image
        last_para = document.add_paragraph()
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = last_para.add_run()
        run.add_picture(image_path, width=Inches(5))
        # Add caption
        caption = document.add_paragraph('图片：{}示意图'.format(topic))
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.style = 'Caption'
    
    # Add table of contents (simplified)
    document.add_page_break()
    toc_heading = document.add_heading('目录', level=1)
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph('一、引言........................................1')
    document.add_paragraph('二、论点一......................................2')
    document.add_paragraph('三、论点二......................................3')
    document.add_paragraph('四、结论........................................4')
    
    # Add references
    document.add_page_break()
    ref_heading = document.add_heading('参考文献', level=1)
    ref_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph('[1] 某某某. 某某研究[J]. 某某期刊, 2020, 10(1): 1-10.')
    document.add_paragraph('[2] 某某某. 某某分析[M]. 某某出版社, 2019.')
    
    # Save document to a temporary location for Vercel
    filename = str(uuid.uuid4()) + '.docx'
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    document.save(filepath)
    return filepath

# This is the handler for Vercel
def vercel_handler(request):
    return app(request.environ, lambda status, headers: None)

if __name__ == '__main__':
    app.run(debug=True)